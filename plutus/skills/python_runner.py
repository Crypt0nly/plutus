"""Python Skill Runner — executes complex Python-based skills.

Unlike simple JSON skills (linear step sequences), Python skills are full
Python scripts that can use loops, conditionals, LLM calls, browser
automation, file I/O, and any Python library.

Every Python skill must define:
    async def run(ctx: PlutusContext, params: dict) -> dict

The PlutusContext provides a clean API for everything a skill needs:
  - Browser automation (Playwright)
  - LLM calls (ask questions, get JSON responses)
  - File I/O (read, write, create documents)
  - Shell commands
  - State persistence (across runs)
  - Logging
"""

from __future__ import annotations

import asyncio
import importlib.util
import json
import logging
import os
import sys
import time
import traceback
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger("plutus.skills.python_runner")

# Where skills store persistent state
STATE_DIR = Path.home() / ".plutus" / "skill_state"
SKILLS_DIR = Path.home() / ".plutus" / "skills"
OUTPUT_DIR = Path.home() / ".plutus" / "skill_output"


class PlutusContext:
    """Everything a Python skill needs to interact with the world.
    
    This is the single object passed to every Python skill's run() function.
    It provides a clean, high-level API so skills don't need to know about
    Plutus internals.
    """

    def __init__(
        self,
        skill_name: str,
        llm_client: Any = None,
        llm_model: str = "gpt-4.1-mini",
    ):
        self.skill_name = skill_name
        self._llm = llm_client
        self._llm_model = llm_model
        self._browser = None
        self._browser_instance = None
        self._browser_context = None
        self._page = None
        self._playwright = None
        self._logs: list[str] = []
        self._state_file = STATE_DIR / f"{skill_name}.json"
        self._output_dir = OUTPUT_DIR / skill_name

        # Ensure directories exist
        STATE_DIR.mkdir(parents=True, exist_ok=True)
        self._output_dir.mkdir(parents=True, exist_ok=True)

    # ── Browser Automation ─────────────────────────────────

    async def browser_open(self):
        """Open a browser and return the Playwright page object.
        
        Returns the raw Playwright Page for full flexibility.
        The browser is headless by default.
        """
        if self._page is not None:
            return self._page

        try:
            from playwright.async_api import async_playwright
            self._playwright = await async_playwright().start()
            
            # Try to reuse existing browser session cookies
            session_dir = Path.home() / ".plutus" / "browser_session"
            session_dir.mkdir(parents=True, exist_ok=True)
            
            self._browser_instance = await self._playwright.chromium.launch(
                headless=True,
                args=["--no-sandbox", "--disable-dev-shm-usage"],
            )
            
            # Try to load saved storage state
            storage_state = session_dir / "state.json"
            if storage_state.exists():
                self._browser_context = await self._browser_instance.new_context(
                    storage_state=str(storage_state),
                )
            else:
                self._browser_context = await self._browser_instance.new_context()
            
            self._page = await self._browser_context.new_page()
            self.log("Browser opened")
            return self._page
        except Exception as e:
            self.log(f"Failed to open browser: {e}")
            raise

    async def browser_navigate(self, url: str, wait_until: str = "domcontentloaded") -> None:
        """Navigate to a URL."""
        page = await self.browser_open()
        await page.goto(url, wait_until=wait_until, timeout=30000)
        await asyncio.sleep(1)
        self.log(f"Navigated to {url}")

    async def browser_click(self, selector: str, timeout: float = 5000) -> None:
        """Click an element by CSS selector."""
        page = await self.browser_open()
        await page.click(selector, timeout=timeout)
        self.log(f"Clicked: {selector}")

    async def browser_type(self, selector: str, text: str, timeout: float = 5000) -> None:
        """Type text into an element."""
        page = await self.browser_open()
        await page.fill(selector, text, timeout=timeout)
        self.log(f"Typed into: {selector}")

    async def browser_evaluate(self, js: str) -> Any:
        """Execute JavaScript in the browser and return the result."""
        page = await self.browser_open()
        return await page.evaluate(js)

    async def browser_get_text(self) -> str:
        """Get all visible text from the current page."""
        page = await self.browser_open()
        return await page.evaluate("() => document.body.innerText")

    async def browser_wait(self, selector: str, timeout: float = 10000) -> None:
        """Wait for an element to appear."""
        page = await self.browser_open()
        await page.wait_for_selector(selector, timeout=timeout)

    async def browser_screenshot(self, name: str = "screenshot") -> str:
        """Take a screenshot and save it. Returns the file path."""
        page = await self.browser_open()
        path = str(self._output_dir / f"{name}.png")
        await page.screenshot(path=path)
        self.log(f"Screenshot saved: {path}")
        return path

    async def browser_upload(self, selector: str, file_path: str) -> None:
        """Upload a file to a file input element."""
        page = await self.browser_open()
        await page.set_input_files(selector, file_path)
        self.log(f"Uploaded {file_path} to {selector}")

    async def browser_save_session(self) -> None:
        """Save the current browser session (cookies, storage) for reuse."""
        if self._browser_context:
            session_dir = Path.home() / ".plutus" / "browser_session"
            session_dir.mkdir(parents=True, exist_ok=True)
            await self._browser_context.storage_state(
                path=str(session_dir / "state.json")
            )
            self.log("Browser session saved")

    async def browser_close(self) -> None:
        """Close the browser."""
        if self._browser_context:
            try:
                await self._browser_context.close()
            except Exception:
                pass
        if self._browser_instance:
            try:
                await self._browser_instance.close()
            except Exception:
                pass
        if hasattr(self, '_playwright') and self._playwright:
            try:
                await self._playwright.stop()
            except Exception:
                pass
        self._page = None
        self._browser_instance = None
        self._browser_context = None
        self.log("Browser closed")

    # ── LLM Calls ──────────────────────────────────────────

    async def llm_ask(self, prompt: str, system: str = "") -> str:
        """Ask the LLM a question and get a text response.
        
        Args:
            prompt: The user message / question
            system: Optional system prompt for context
            
        Returns:
            The LLM's text response
        """
        if self._llm is None:
            self._llm = self._create_llm_client()

        messages = []
        if system:
            messages.append({"role": "system", "content": system})
        messages.append({"role": "user", "content": prompt})

        try:
            response = await self._llm.chat.completions.create(
                model=self._llm_model,
                messages=messages,
                max_tokens=4096,
            )
            result = response.choices[0].message.content or ""
            self.log(f"LLM response: {len(result)} chars")
            return result
        except Exception as e:
            self.log(f"LLM call failed: {e}")
            raise

    async def llm_json(self, prompt: str, system: str = "") -> dict:
        """Ask the LLM and parse the response as JSON.
        
        The prompt should instruct the LLM to respond in JSON format.
        """
        response = await self.llm_ask(prompt, system)
        
        # Try to extract JSON from the response
        text = response.strip()
        
        # Handle markdown code blocks
        if "```json" in text:
            text = text.split("```json")[1].split("```")[0].strip()
        elif "```" in text:
            text = text.split("```")[1].split("```")[0].strip()
        
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            self.log(f"Failed to parse LLM response as JSON: {text[:200]}")
            return {"raw_response": response, "parse_error": True}

    def _create_llm_client(self):
        """Create an OpenAI-compatible LLM client."""
        from openai import AsyncOpenAI
        return AsyncOpenAI()  # Uses OPENAI_API_KEY env var

    # ── File I/O ───────────────────────────────────────────

    def write_file(self, path: str, content: str) -> str:
        """Write content to a file. Creates directories if needed.
        
        If path is relative, it's relative to the skill's output directory.
        Returns the absolute path.
        """
        if not os.path.isabs(path):
            path = str(self._output_dir / path)
        
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        Path(path).write_text(content, encoding="utf-8")
        self.log(f"Wrote file: {path} ({len(content)} chars)")
        return path

    def read_file(self, path: str) -> str:
        """Read a file and return its content."""
        if not os.path.isabs(path):
            path = str(self._output_dir / path)
        return Path(path).read_text(encoding="utf-8")

    def create_document(
        self,
        title: str,
        content: str,
        format: str = "txt",
        filename: str = "",
    ) -> str:
        """Create a document file (txt, md, or docx).
        
        Returns the file path.
        """
        if not filename:
            safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)
            filename = f"{safe_title}.{format}"
        
        path = str(self._output_dir / filename)

        if format == "docx":
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(title, 0)
                for paragraph in content.split("\n\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
                doc.save(path)
            except ImportError:
                # Fallback to txt
                path = path.replace(".docx", ".txt")
                Path(path).write_text(f"{title}\n{'='*len(title)}\n\n{content}", encoding="utf-8")
        else:
            if format == "md":
                full_content = f"# {title}\n\n{content}"
            else:
                full_content = f"{title}\n{'='*len(title)}\n\n{content}"
            Path(path).write_text(full_content, encoding="utf-8")

        self.log(f"Created document: {path}")
        return path

    # ── Shell Commands ─────────────────────────────────────

    async def shell(self, command: str, timeout: float = 30) -> str:
        """Run a shell command and return the output."""
        try:
            proc = await asyncio.create_subprocess_shell(
                command,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await asyncio.wait_for(
                proc.communicate(), timeout=timeout
            )
            output = stdout.decode("utf-8", errors="replace")
            if stderr:
                output += "\n" + stderr.decode("utf-8", errors="replace")
            self.log(f"Shell: {command[:50]}... → {len(output)} chars")
            return output.strip()
        except asyncio.TimeoutError:
            proc.kill()
            return f"[TIMEOUT] Command timed out after {timeout}s"
        except Exception as e:
            return f"[ERROR] {e}"

    # ── State Persistence ──────────────────────────────────

    def save_state(self, key: str, value: Any) -> None:
        """Save a value to persistent state (survives restarts).
        
        State is stored per-skill in ~/.plutus/skill_state/{skill_name}.json
        """
        state = self._load_full_state()
        state[key] = value
        self._state_file.write_text(json.dumps(state, indent=2, default=str))

    def load_state(self, key: str, default: Any = None) -> Any:
        """Load a value from persistent state."""
        state = self._load_full_state()
        return state.get(key, default)

    def _load_full_state(self) -> dict:
        if self._state_file.exists():
            try:
                return json.loads(self._state_file.read_text())
            except Exception:
                return {}
        return {}

    def clear_state(self) -> None:
        """Clear all persistent state for this skill."""
        if self._state_file.exists():
            self._state_file.unlink()

    # ── Logging ────────────────────────────────────────────

    def log(self, message: str) -> None:
        """Log a message (visible in skill execution results)."""
        timestamp = time.strftime("%H:%M:%S")
        entry = f"[{timestamp}] {message}"
        self._logs.append(entry)
        logger.info(f"[{self.skill_name}] {message}")

    def get_logs(self) -> list[str]:
        """Get all log entries."""
        return self._logs.copy()

    # ── Cleanup ────────────────────────────────────────────

    async def cleanup(self) -> None:
        """Clean up resources (called automatically after skill execution)."""
        await self.browser_close()


@dataclass
class PythonSkillResult:
    """Result of a Python skill execution."""
    success: bool
    skill_name: str
    result: Any = None
    error: str | None = None
    logs: list[str] = field(default_factory=list)
    duration: float = 0.0
    output_files: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "success": self.success,
            "skill_name": self.skill_name,
            "result": self.result,
            "error": self.error,
            "logs": self.logs,
            "duration_seconds": round(self.duration, 2),
            "output_files": self.output_files,
        }


class PythonSkillRunner:
    """Loads and executes Python-based skills."""

    def __init__(self, llm_model: str = "gpt-4.1-mini"):
        self._llm_model = llm_model

    async def run(
        self,
        skill_name: str,
        params: dict[str, Any],
        script_path: str | None = None,
    ) -> PythonSkillResult:
        """Execute a Python skill.
        
        Args:
            skill_name: Name of the skill
            params: Parameters to pass to the skill's run() function
            script_path: Path to the .py file (if None, looks in SKILLS_DIR)
        """
        start_time = time.time()

        # Find the script
        if script_path is None:
            script_path = str(SKILLS_DIR / f"{skill_name}.py")

        if not os.path.exists(script_path):
            return PythonSkillResult(
                success=False,
                skill_name=skill_name,
                error=f"Script not found: {script_path}",
            )

        # Create the context
        ctx = PlutusContext(
            skill_name=skill_name,
            llm_model=self._llm_model,
        )

        try:
            # Load the skill module dynamically
            spec = importlib.util.spec_from_file_location(
                f"plutus_skill_{skill_name}", script_path
            )
            if spec is None or spec.loader is None:
                return PythonSkillResult(
                    success=False,
                    skill_name=skill_name,
                    error=f"Failed to load module from {script_path}",
                )

            module = importlib.util.module_from_spec(spec)
            
            # Add the skills directory to sys.path so skills can import each other
            skills_dir = str(Path(script_path).parent)
            if skills_dir not in sys.path:
                sys.path.insert(0, skills_dir)

            spec.loader.exec_module(module)

            # Find the run function
            if not hasattr(module, "run"):
                return PythonSkillResult(
                    success=False,
                    skill_name=skill_name,
                    error="Skill must define 'async def run(ctx, params) -> dict'",
                )

            run_func = module.run
            if not asyncio.iscoroutinefunction(run_func):
                return PythonSkillResult(
                    success=False,
                    skill_name=skill_name,
                    error="run() must be an async function (async def run(ctx, params))",
                )

            # Execute the skill with a timeout to prevent infinite loops
            # from hanging the backend indefinitely.
            SKILL_TIMEOUT = 300  # 5 minutes max
            ctx.log(f"Starting skill: {skill_name}")
            try:
                result = await asyncio.wait_for(run_func(ctx, params), timeout=SKILL_TIMEOUT)
            except asyncio.TimeoutError:
                ctx.log(f"TIMEOUT: Skill exceeded {SKILL_TIMEOUT}s limit")
                return PythonSkillResult(
                    success=False,
                    skill_name=skill_name,
                    error=f"Skill timed out after {SKILL_TIMEOUT} seconds",
                    logs=ctx.get_logs(),
                    duration=time.time() - start_time,
                )
            except SystemExit as e:
                ctx.log(f"Skill called sys.exit({e.code})")
                return PythonSkillResult(
                    success=False,
                    skill_name=skill_name,
                    error=f"Skill called sys.exit({e.code}) — blocked to protect the backend",
                    logs=ctx.get_logs(),
                    duration=time.time() - start_time,
                )

            # Validate the result
            if not isinstance(result, dict):
                result = {"result": str(result)}

            success = result.get("success", True)

            # Collect output files
            output_dir = OUTPUT_DIR / skill_name
            output_files = []
            if output_dir.exists():
                for f in output_dir.iterdir():
                    if f.is_file():
                        output_files.append(str(f))

            return PythonSkillResult(
                success=success,
                skill_name=skill_name,
                result=result,
                logs=ctx.get_logs(),
                duration=time.time() - start_time,
                output_files=output_files,
            )

        except Exception as e:
            tb = traceback.format_exc()
            ctx.log(f"FATAL ERROR: {e}")
            logger.error(f"Python skill {skill_name} failed:\n{tb}")
            return PythonSkillResult(
                success=False,
                skill_name=skill_name,
                error=f"{type(e).__name__}: {e}",
                logs=ctx.get_logs(),
                duration=time.time() - start_time,
            )

        finally:
            await ctx.cleanup()
