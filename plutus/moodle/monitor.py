"""
MoodleMonitor — Orchestrates assignment detection, completion, and submission.

This is the main entry point. It:
1. Logs into Moodle
2. Scans all courses for assignments
3. Identifies which ones need action (unsubmitted, not yet graded)
4. Uses the LLM to generate responses based on course materials
5. Submits the completed work
6. Reports results
"""

from __future__ import annotations

import asyncio
import json
import logging
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from typing import Optional

from plutus.moodle.client import MoodleClient, Assignment, Course

logger = logging.getLogger("plutus.moodle.monitor")


class MoodleMonitor:
    """
    Full assignment monitoring pipeline.
    
    Usage:
        monitor = MoodleMonitor(
            email="student@unyp.cz",
            password="password123",
            base_url="https://elearning.unyp.cz",
        )
        results = await monitor.run()
    """

    def __init__(
        self,
        email: str,
        password: str,
        base_url: str = "https://elearning.unyp.cz",
        auto_submit: bool = False,
        course_filter: Optional[list[str]] = None,
    ):
        self.email = email
        self.password = password
        self.base_url = base_url
        self.auto_submit = auto_submit
        self.course_filter = course_filter  # Only check these courses (by name substring)
        self.client = MoodleClient(base_url)
        self._output_dir = Path.home() / ".plutus" / "moodle_assignments"
        self._output_dir.mkdir(parents=True, exist_ok=True)
        self._log_path = self._output_dir / "monitor_log.json"

    async def run(self) -> dict:
        """
        Run the full monitoring pipeline.
        
        Returns a summary dict with:
        - courses_checked: number of courses scanned
        - assignments_found: total assignments found
        - pending_assignments: assignments needing action
        - completed: assignments that were completed and submitted
        - errors: any errors encountered
        """
        results = {
            "timestamp": datetime.now().isoformat(),
            "courses_checked": 0,
            "assignments_found": 0,
            "pending_assignments": [],
            "completed": [],
            "errors": [],
        }

        try:
            # Step 1: Login
            logger.info("Logging into Moodle...")
            if not await self.client.login(self.email, self.password):
                results["errors"].append("Login failed")
                return results

            # Step 2: Get courses
            logger.info("Fetching courses...")
            courses = await self.client.get_courses()
            
            # Apply course filter if specified
            if self.course_filter:
                courses = [
                    c for c in courses
                    if any(f.lower() in c.name.lower() for f in self.course_filter)
                ]
            
            results["courses_checked"] = len(courses)
            logger.info(f"Checking {len(courses)} courses")

            # Step 3: Scan each course for assignments
            all_assignments = []
            for course in courses:
                logger.info(f"Scanning: {course.name}")
                assignments = await self.client.get_assignments(course)
                
                # Get details for each assignment
                for assignment in assignments:
                    assignment = await self.client.get_assignment_details(assignment)
                    all_assignments.append(assignment)
                    await asyncio.sleep(0.5)  # Be gentle with the server

            results["assignments_found"] = len(all_assignments)

            # Step 4: Identify pending assignments
            pending = [a for a in all_assignments if a.needs_action]
            results["pending_assignments"] = [
                {
                    "title": a.title,
                    "course": a.course_name,
                    "due": a.due,
                    "status": a.status,
                    "is_overdue": a.is_overdue,
                    "url": a.url,
                    "description": a.description[:200] if a.description else "",
                }
                for a in pending
            ]

            logger.info(f"Found {len(pending)} assignments needing action")

            # Step 5: Complete and submit if auto_submit is enabled
            if self.auto_submit and pending:
                for assignment in pending:
                    try:
                        result = await self._complete_assignment(assignment)
                        results["completed"].append(result)
                    except Exception as e:
                        results["errors"].append(f"Failed to complete {assignment.title}: {e}")

            # Save the log
            self._save_log(results)

            return results

        except Exception as e:
            logger.error(f"Monitor run failed: {e}")
            results["errors"].append(str(e))
            return results

        finally:
            await self.client.close()

    async def check_only(self) -> list[dict]:
        """
        Just check for pending assignments without completing them.
        
        Returns a list of pending assignment summaries.
        """
        try:
            if not await self.client.login(self.email, self.password):
                return [{"error": "Login failed"}]

            courses = await self.client.get_courses()
            if self.course_filter:
                courses = [
                    c for c in courses
                    if any(f.lower() in c.name.lower() for f in self.course_filter)
                ]

            pending = []
            for course in courses:
                assignments = await self.client.get_assignments(course)
                for assignment in assignments:
                    assignment = await self.client.get_assignment_details(assignment)
                    if assignment.needs_action:
                        pending.append({
                            "title": assignment.title,
                            "course": assignment.course_name,
                            "due": assignment.due,
                            "status": assignment.status,
                            "is_overdue": assignment.is_overdue,
                            "url": assignment.url,
                            "description": assignment.description[:300] if assignment.description else "",
                        })
                    await asyncio.sleep(0.5)

            return pending

        finally:
            await self.client.close()

    async def _complete_assignment(self, assignment: Assignment) -> dict:
        """
        Use the LLM to complete an assignment and submit it.
        
        This is the core intelligence — it:
        1. Reads the assignment instructions
        2. Gets relevant course materials for context
        3. Generates a response using the LLM
        4. Creates a document (Word/PDF)
        5. Submits it to Moodle
        """
        logger.info(f"Completing: {assignment.title} ({assignment.course_name})")

        try:
            # Get course materials for context
            course = Course(
                id=assignment.course_id,
                name=assignment.course_name,
                url=f"{self.base_url}/course/view.php?id={assignment.course_id}",
            )
            
            # Try to determine the week number from the assignment title
            import re
            week_match = re.search(r'Week\s*(\d+)', assignment.title, re.IGNORECASE)
            week = int(week_match.group(1)) if week_match else None
            
            materials = await self.client.get_course_materials(course, week)

            # Build the prompt for the LLM
            prompt = self._build_completion_prompt(assignment, materials)

            # Generate the response using Plutus's LLM client
            response_text = await self._generate_response(prompt)

            if not response_text:
                return {
                    "assignment": assignment.title,
                    "success": False,
                    "error": "LLM failed to generate response",
                }

            # Save the response as a document
            doc_path = await self._create_document(assignment, response_text)

            # Submit to Moodle
            if doc_path:
                submit_result = await self.client.submit_file(assignment, str(doc_path))
                return {
                    "assignment": assignment.title,
                    "course": assignment.course_name,
                    "success": submit_result.get("success", False),
                    "document": str(doc_path),
                    "error": submit_result.get("error", ""),
                }
            else:
                # Try online text submission
                submit_result = await self.client.submit_online_text(assignment, response_text)
                return {
                    "assignment": assignment.title,
                    "course": assignment.course_name,
                    "success": submit_result.get("success", False),
                    "error": submit_result.get("error", ""),
                }

        except Exception as e:
            return {
                "assignment": assignment.title,
                "success": False,
                "error": str(e),
            }

    def _build_completion_prompt(self, assignment: Assignment, materials: list[dict]) -> str:
        """Build a prompt for the LLM to complete the assignment."""
        materials_text = ""
        if materials:
            materials_text = "\n\nAvailable course materials for context:\n"
            for m in materials[:10]:  # Limit to 10 materials
                materials_text += f"- {m['title']} ({m['type']})\n"

        prompt = f"""You are a university student completing an assignment. Write a thorough, 
well-researched response that demonstrates understanding of the course material.

Course: {assignment.course_name}
Assignment: {assignment.title}
Instructions: {assignment.description or 'No specific instructions provided. Write a thoughtful response based on the topic.'}
{materials_text}

Requirements:
1. Write in academic style with proper paragraphs (not bullet points)
2. Include relevant examples and analysis
3. Reference course concepts and theories discussed in class
4. Length: 500-1000 words for weekly assignments, 2000-3000 for major papers
5. Use APA format for any citations
6. Be original and demonstrate critical thinking

Write the complete assignment response:"""

        return prompt

    async def _generate_response(self, prompt: str) -> Optional[str]:
        """Generate an assignment response using the LLM."""
        try:
            from plutus.core.llm import LLMClient
            from plutus.config import load_config

            config = load_config()
            llm = LLMClient(config)

            messages = [
                {"role": "system", "content": "You are a diligent university student writing an assignment. Write thorough, well-structured academic responses."},
                {"role": "user", "content": prompt},
            ]

            response = await llm.chat(messages)
            
            if response and response.get("content"):
                return response["content"]
            return None

        except Exception as e:
            logger.error(f"LLM generation failed: {e}")
            return None

    async def _create_document(self, assignment: Assignment, text: str) -> Optional[Path]:
        """Create a Word document from the response text."""
        try:
            # Try to create a .docx file
            try:
                from docx import Document
                
                doc = Document()
                doc.add_heading(assignment.title, level=1)
                doc.add_paragraph(f"Course: {assignment.course_name}")
                doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
                doc.add_paragraph("")  # Blank line
                
                # Split text into paragraphs and add them
                paragraphs = text.split("\n\n")
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                
                # Save
                safe_title = re.sub(r'[^\w\s-]', '', assignment.title).strip()
                safe_title = re.sub(r'\s+', '_', safe_title)
                filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d')}.docx"
                filepath = self._output_dir / filename
                doc.save(str(filepath))
                
                logger.info(f"Created document: {filepath}")
                return filepath
                
            except ImportError:
                # Fallback: create a plain text file
                import re as re_mod
                safe_title = re_mod.sub(r'[^\w\s-]', '', assignment.title).strip()
                safe_title = re_mod.sub(r'\s+', '_', safe_title)
                filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d')}.txt"
                filepath = self._output_dir / filename
                
                with open(filepath, 'w') as f:
                    f.write(f"{assignment.title}\n")
                    f.write(f"Course: {assignment.course_name}\n")
                    f.write(f"Date: {datetime.now().strftime('%B %d, %Y')}\n")
                    f.write(f"\n{'='*60}\n\n")
                    f.write(text)
                
                logger.info(f"Created text file: {filepath}")
                return filepath

        except Exception as e:
            logger.error(f"Failed to create document: {e}")
            return None

    def _save_log(self, results: dict):
        """Save monitoring results to a log file."""
        try:
            # Load existing log
            if self._log_path.exists():
                log = json.loads(self._log_path.read_text())
            else:
                log = {"runs": []}

            log["runs"].append(results)
            
            # Keep last 50 runs
            log["runs"] = log["runs"][-50:]
            
            self._log_path.write_text(json.dumps(log, indent=2, default=str))
        except Exception as e:
            logger.warning(f"Failed to save log: {e}")

    async def get_status(self) -> dict:
        """Get the status of the last monitoring run."""
        try:
            if self._log_path.exists():
                log = json.loads(self._log_path.read_text())
                if log.get("runs"):
                    return log["runs"][-1]
            return {"status": "No runs yet"}
        except Exception:
            return {"status": "Error reading log"}
